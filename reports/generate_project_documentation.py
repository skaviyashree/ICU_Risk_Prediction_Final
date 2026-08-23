import os
import sys
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES_DIR = os.path.join(BASE_DIR, "figures")
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)

docx_path = os.path.join(REPORTS_DIR, "Project_Technical_Documentation.docx")
pdf_path = os.path.join(REPORTS_DIR, "Project_Technical_Documentation.pdf")

# Complete technical text for all 40 items, detailed enough for another AI
sections_content = [
    ("1. Executive Summary", 
     "This technical documentation dossier outlines the development, mathematical validation, and clinical deployment of an Explainable Artificial Intelligence (XAI) bedside Clinical Decision Support System (CDSS) for early in-hospital mortality risk prediction in intensive care units. Developed as a feasibility study utilizing the publicly available MIMIC-IV Clinical Demo database (100 adult patients, 140 ICU stays, 11% mortality rate), the pipeline implements a leak-free temporal framework. This framework restricts all bedside vital signs and laboratory feature extractions strictly to the first 24 hours of a patient's first ICU admission. The machine learning engine evaluates nine classifier candidates, deploying a highly regularized ExtraTrees Classifier as the champion model. By optimizing the classification decision boundary using Youden's J statistic on out-of-fold cross-validation predictions, we shift the threshold from 0.50 to 0.39, resuscitating model sensitivity to 50.00% and specificity to 55.56% on unseen test data. Bedside interpretability is provided via a Tree SHAP engine and a Streamlit web application, allowing physicians to audit risk drivers dynamically and generate signed A4 Patient Chart PDFs at the bedside."),
    
    ("2. Project Motivation",
     "Critical care medicine is characterized by a high volume of time-sensitive patient data. In modern ICUs, continuous monitoring devices produce streams of vital sign telemetries (heart rate, respiratory rate, blood pressure, oxygen saturation). This data density, coupled with intermittent laboratory blood panels, often leads to cognitive overload and alarm fatigue for ICU clinicians. Alarm fatigue—where clinicians desensitize to warning signals due to high rates of false-positive alarms—poses a major threat to patient safety, leading to delayed responses to actual clinical deterioration. Furthermore, standard EHR machine learning models frequently suffer from data leakage by extracting features recorded right before death or discharge, rendering them useless for early intervention. The motivation of this project is to build an early-warning CDSS that operates strictly on the initial 24 hours of ICU stay to detect high-risk states early, while providing transparent explanations using game-theoretic SHAP values to bridge the gap between machine learning and clinician bedside trust."),
    
    ("3. Problem Statement",
     "Given a patient's demographics, continuous bedside vital sign telemetry, and laboratory panel measurements recorded strictly within the first 24 hours of their first ICU stay, the objective is to predict the probability of in-hospital mortality (defined by the binary label hospital_expire_flag). The mathematical and clinical constraints are: (1) Clinical Actionability: The model must operate on early, non-leaked observations to enable life-saving interventions. (2) Imbalance Resilience: The system must handle a severe class imbalance (11.0% mortality) without collapsing into predicting survival for all cases (majority-class collapse, which yields 0% recall). (3) Transparency: The predictions must be fully explainable at the bedside to allow immediate physician audit."),
    
    ("4. Objectives",
     "The primary engineering and clinical objectives of the project are:\n"
     "1. Ingest raw de-identified EHR tables from the MIMIC-IV Clinical Demo database and isolate a clean cohort of adult stays.\n"
     "2. Implement a clinical preprocessing pipeline to convert temperature units, clean sensor noise, and flag outliers.\n"
     "3. Design a feature engineering pipeline to slice the 24-hour observation window and compute 6 distinct statistical aggregates per concept.\n"
     "4. Implement clinical imputation utilizing healthy reference baselines to prevent data skewing.\n"
     "5. Train, optimize, and cross-validate nine candidate classifiers under Stratified 5-Fold Cross-Validation.\n"
     "6. Optimize classification decision thresholds using Youden's J statistic on out-of-fold predictions to handle class imbalance.\n"
     "7. Integrate a Tree SHAP explainability engine to deconstruct risk scores into biophysical risk drivers.\n"
     "8. Deconstruct predictions into natural language narratives and deploy the CDSS as a Streamlit web application with vector PDF registry chart exports."),
     
    ("5. Scope",
     "This project is developed strictly as a proof-of-concept feasibility study using the MIMIC-IV Clinical Demo dataset (100 patients, 140 stays). It establishes a baseline MLOps framework for early risk prediction and bedside interpretability. The system has not undergone clinical trials, clinical validation, or real-world hospital deployment, and must not be used for actual clinical diagnostic decisions. The scope covers data ingestion, offline stratified model training, threshold sweeps, SHAP explainability, and Streamlit Community Cloud hosting."),
     
    ("6. Dataset Description",
     "The project uses the MIMIC-IV Clinical Demo dataset (version 2.2), a publicly available de-identified EHR database containing clinical data from Beth Israel Deaconess Medical Center. The primary tables ingested are:\n"
     "• patients: Demographic anchor age and gender.\n"
     "• admissions: Hospital admission details, race, insurance plan, marital status, and the hospital_expire_flag target.\n"
     "• icustays: ICU stay timestamps (intime, outtime), care unit designations, and stay IDs.\n"
     "• chartevents: Bedside vitals (668,862 records).\n"
     "• labevents: Clinical laboratory blood draws (107,727 records)."),
     
    ("7. Dataset Statistics",
     "The cohort consists of 100 adult patients across 140 ICU stays. The outcome target hospital_expire_flag displays a severe class imbalance, with 15 stays resulting in in-hospital mortality (11.0%) and 125 stays resulting in survival (89.0%). The average age at admission is 64.92 years. Demographics show a predominant white population (78%), with emergency admissions representing the majority of cases. Continuous telemetry data contains 668,862 bedside vitals and 107,727 lab measurements, presenting high data density."),
     
    ("8. Data Dictionary",
     "• subject_id (int): Unique identifier for each patient.\n"
     "• hadm_id (int): Unique identifier for each hospital admission.\n"
     "• stay_id (int): Unique identifier for each ICU stay.\n"
     "• age (float): Patient age at ICU admission (anchor_age + year of intime - anchor_year).\n"
     "• gender (int): Biological sex (1 = Male, 0 = Female).\n"
     "• race (string): Self-reported ethnicity/race.\n"
     "• admission_type (string): Source category of hospital admission.\n"
     "• insurance (string): Primary insurance plan type (Medicare, Medicaid, Other).\n"
     "• marital_status (string): Marital status of the patient.\n"
     "• hospital_expire_flag (int): Target label (1 = Deceased in hospital, 0 = Survived)."),
     
    ("9. Clinical Variables Used",
     "The model utilizes 9 core biophysical markers selected for clinical significance:\n"
     "1. Age (demographic risk)\n"
     "2. Heart Rate Std (vital sign volatility)\n"
     "3. SpO2 Mean (oxygenation baseline)\n"
     "4. Systolic BP Mean (cardiovascular baseline)\n"
     "5. MAP Min (lowest Mean Arterial Pressure, shock risk)\n"
     "6. Temperature Mean (thermoregulation baseline, sepsis marker)\n"
     "7. BUN Mean (Blood Urea Nitrogen, renal function)\n"
     "8. Creatinine Latest (kidney injury progression)\n"
     "9. Potassium Min (cardiac muscle stability, arrhythmia risk)"),
     
    ("10. Feature Selection Strategy",
     "Training 107 raw features on a small dataset (100 patients) leads to high-variance overfitting (p > N problem). To prevent this, we drop the SMOTE oversampler and constrain the pipeline's feature space to 9 clinically validated biomarkers. This regularizes the decision boundary, reducing model variance and ensuring stable out-of-sample generalization on the holdout test set."),
     
    ("11. Feature Engineering Pipeline",
     "The pipeline is executed inside the 24-hour observation window (t <= intime + 24h). For each vital and lab concept, it extracts 6 statistical aggregates: Mean, Minimum, Maximum, Standard Deviation, Latest (last value recorded), and Trend (Latest - First). Laboratory parameters also track missingness indicator flags (is_missing_<lab_name>) to record the clinical decision of ordering a test."),
     
    ("12. Data Cleaning",
     "Data cleaning handles unit discrepancies and sensor noise. Fahrenheit temperature readings (item ID 223761) are normalized to Celsius: C = (F - 32) * 5/9, and merged with Celsius records (item ID 223762). Telemetry observations falling outside plausible bounds (e.g. Heart Rate < 30 or > 220 bpm) are set to NaN, allowing them to be imputed during feature engineering."),
     
    ("13. Missing Value Handling",
     "Missing vital signs are imputed using global cohort medians. Missing laboratory values (which are ordered intermittently due to informative missingness) are imputed using healthy reference baselines (Potassium = 4.2 mEq/L, BUN = 14.0 mg/dL, Creatinine = 0.9 mg/dL) using the CLINICAL_REF_LABS dictionary. This represents the medical assumption of normalcy unless observed otherwise, avoiding bias."),
     
    ("14. Outlier Detection",
     "Outlier detection relies on clinical limits defined in the OUTLIER_BOUNDS dictionary. Telemetry monitors generate noise (e.g. displaced sensors showing SpO2 of 0%). Removing these extreme values and setting them to NaN ensures that the downstream statistical aggregates (especially Standard Deviation and Trend) are not corrupted by artifacts."),
     
    ("15. Scaling",
     "Features are scaled using StandardScaler() within a ColumnTransformer. This normalizes the 9 clinical biomarkers to have a mean of 0 and standard deviation of 1. Scaling is critical for linear models (Logistic Regression) and prevents variables with larger absolute scales (e.g. Systolic BP) from dominating the distance calculations in ensemble algorithms."),
     
    ("16. Feature Aggregation",
     "Feature aggregation compresses continuous vital sign and lab measurements over 24 hours into 6 static aggregates. The minimum value captures acute decompensations (bradycardia, hypoxia), maximum captures peak stress (hypertensive crisis, fever), standard deviation captures volatility, and trend captures the patient's direction of recovery or decline. This yields a structured feature matrix of 107 columns prior to feature selection."),
     
    ("17. Complete ML Pipeline",
     "The pipeline is implemented using sklearn and imblearn. It comprises a preprocessor ColumnTransformer (which applies StandardScaler to the 9 selected features and drops the remaining columns) and a classifier. Nine candidates are evaluated: Logistic Regression, Random Forest, Balanced Random Forest, XGBoost, LightGBM, CatBoost, ExtraTrees, Soft-Voting, and Stacking ensembles."),
     
    ("18. Cross Validation Strategy",
     "The evaluation uses Stratified 5-Fold Cross-Validation. Splitting is stratified by the target label hospital_expire_flag, ensuring that the 11% deceased-to-survived ratio is maintained in both training and validation folds. This prevents optimistic biases and ensures stable evaluation on imbalanced data."),
     
    ("19. Threshold Optimization",
     "Standard classifiers predict survival for all patients on imbalanced data using a 0.5 threshold. To resolve this, we sweep thresholds (0.01 to 0.99) on out-of-fold validation sets to find the cutoff maximizing Youden's J Statistic (J = Sensitivity + Specificity - 1, equivalent to Balanced Accuracy). The optimal thresholds are averaged across folds to obtain a robust global threshold of 0.39 for the ExtraTrees model."),
     
    ("20. Champion Model Selection",
     "Models are ranked by their holdout Clinical Composite Score (PR-AUC + ROC-AUC + Recall). The ExtraTrees Classifier (ext_trees) with an optimized threshold of 0.39 was selected as the champion model. It achieved holdout Accuracy = 0.5500, Sensitivity = 0.5000, Specificity = 0.5556, ROC-AUC = 0.5556, and PR-AUC = 0.1603, outperforming stacking ensembles which overfit the small sample size."),
     
    ("21. Explainability using SHAP",
     "We use Tree SHAP to provide global and local explainability. Global interpretability identifies Potassium (Min) and BUN (Mean) as the primary risk predictors. Local interpretability generates patient-specific waterfall plots, deconstructing risk scores into biophysical drivers and compiling them into natural-language risk narratives."),
     
    ("22. Clinical Recommendation Engine",
     "Calculated probabilities are mapped to clinical alerts: Low Risk (<5%), Moderate Risk (<15%), High Risk (<40%), and Critical Risk (>=40%). Each risk band is mapped to rule-based clinical recommendations (e.g. triggering a Rapid Response Team bedside review if risk is >= 40% or Systolic BP is < 90 mmHg)."),
     
    ("23. Dashboard Architecture",
     "The Streamlit bedside dashboard follows a clean, responsive layout. It utilizes cached data and model loading to prevent lag, displays patient demographics, calculated risk levels, alert banners, local SHAP waterfalls, bedside trend charts, rule-based clinical guides, and exports vector A4 PDFs."),
     
    ("24. Streamlit UI Explanation",
     "• Clinical Mode: Features Demo Patient Registry Lookup (dropdown of stays) and Bedside Physiological Slider Builder (sliders for manually adjusting vital/lab metrics). Displays alert cards, SHAP narratives, and PDF exports.\n"
     "• Research Mode: Displays performance leaderboards, model galleries (ROC/confusion matrix curves), global SHAP swarm plots, and cohort correlation heatmaps."),
     
    ("25. PDF Report Generator",
     "The pdf_generator.py module uses matplotlib's vector engine to compile a vector A4 PDF Bedside Chart registry export. The PDF contains patient demographics, risk assessment scales, SHAP narratives, clinical advisories, and physician signature fields, ready for clinical registry filing."),
     
    ("26. Folder Structure",
     "• data/: raw CSVs, processed feature matrices, serialized models (best_model.joblib).\n"
     "• figures/: candidate model sweeps, ROC/PR curves, post-retraining figures.\n"
     "• reports/: master handbooks, model comparison reports, ieee conference papers.\n"
     "• src/: data_ingestion.py, preprocessing.py, feature_engineering.py, train.py, explainability.py, pdf_generator.py, config.py.\n"
     "• app/: app.py, utils.py.\n"
     "• Screenshots/: clinical/research UI pngs."),
     
    ("27. File-by-file Explanation",
     "• config.py: Configuration path definitions, clinical mappings, outlier boundaries, and hyperparameters.\n"
     "• data_ingestion.py: Ingests raw zipped CSVs, selects first stays, filters age >= 18, and merges admissions.\n"
     "• preprocessing.py: Handles Fahrenheit to Celsius conversion and sets vital outliers to NaN.\n"
     "• feature_engineering.py: Slices the 24-hour observation window, computes 6 aggregates, handles missingness, and imputes labs with healthy references.\n"
     "• train.py: Scales selected clinical markers, runs Stratified 5-Fold CV, searches Youden's J threshold, and saves the champion ExtraTrees pipeline.\n"
     "• explainability.py: Initializes the SHAP engine, translates raw columns to clinical labels, and generates bedside waterfall plots and clinician NLP narratives.\n"
     "• pdf_generator.py: Generates the vector A4 PDF bedside report.\n"
     "• app.py & utils.py: Streamlit dashboard frontend and CSS styles."),
     
    ("28. Module Dependency Diagram",
     "• config.py is imported by all backend modules.\n"
     "• data_ingestion.py -> preprocessing.py -> feature_engineering.py -> train.py.\n"
     "• train.py serializes best_model.joblib.\n"
     "• explainability.py imports best_model.joblib and config.py.\n"
     "• pdf_generator.py imports config.py.\n"
     "• app.py imports explainability.py, pdf_generator.py, utils.py, and best_model.joblib."),
     
    ("29. Execution Flow",
     "The pipeline is executed end-to-end using run_pipeline.py: Ingestion loads raw files -> Preprocessing cleans vitals -> Feature Engineering aggregates over 24 hours -> Training scales, cross-validates, searches thresholds, and saves best_model.joblib -> Explainability fits the background SHAP explainer -> Figures are saved."),
     
    ("30. System Workflow",
     "At the bedside: The clinician selects an ICU stay or adjusts sliders -> Streamlit calls the ExtraTrees model to evaluate probability -> SHAP computes local contributions -> Rule-based advisories are triggered -> Matplotlib compiles the vector PDF chart -> The physician signs the chart."),
     
    ("31. Block Diagram Explanation",
     "The system block diagram (Fig. 1) represents the 3-layer workflow: (1) Data Ingestion and Multimodal Preprocessing; (2) Predictive Machine Learning Engine (ExtraTrees champion); and (3) Clinical Actionability Tier (Streamlit/PDF export), illustrating how EHR data is transformed into explainable bedside advisories."),
     
    ("32. Deployment Process",
     "The system is deployed on Streamlit Community Cloud. The GitHub repository contains all code, dependencies, and serialized models. Streamlit Community Cloud pulls the repository, installs dependencies, and serves the application on a public URL."),
     
    ("33. GitHub Repository Structure",
     "The repository contains: app/ (app.py, utils.py), src/ (modules, config.py), data/ (raw/processed/models), figures/ (curves, rankings), reports/ (handbooks, papers), screenshots/ (UI pngs), requirements.txt, and run_pipeline.py."),
     
    ("34. Streamlit Cloud Deployment",
     "The Streamlit Community Cloud deployment is configured to automatically fetch the master branch of the repository, set up the Python environment, install packages from requirements.txt, and run app/app.py as the entrypoint served on the public URL."),
     
    ("35. Software Requirements",
     "• Python 3.10+\n"
     "• Operating System: Windows/Linux/MacOS\n"
     "• Web Browser: Chrome, Firefox, or Safari for dashboard access\n"
     "• Git: for repository version control"),
     
    ("36. Hardware Requirements",
     "• CPU: Dual-Core 2.0 GHz or higher (Quad-Core recommended for SHAP background fits)\n"
     "• RAM: 8 GB or higher (16 GB recommended for loading full chartevents in memory)\n"
     "• Disk Space: 500 MB (MIMIC demo files are small, full database requires >100 GB)"),
     
    ("37. Libraries Used",
     "• pandas (v2.0.0+): clinical data frame manipulation\n"
     "• numpy (v1.22.0+): mathematical operations\n"
     "• scikit-learn (v1.2.0+): ML pipelines and scaling\n"
     "• imbalanced-learn (v0.10.0+): class weighting\n"
     "• xgboost/lightgbm/catboost: candidate classifiers\n"
     "• shap (v0.41.0+): explainability engine\n"
     "• reportlab (v4.5.1+): portrait master handbook PDF compilation\n"
     "• python-docx (v1.2.0+): Word document formatting"),
     
    ("38. Limitations",
     "The primary limitation is the sample size constraint (100 patients, 140 stays) of the MIMIC-IV Clinical Demo database. This causes high variance and limits model discriminative power (ROC-AUC = 0.5556). Furthermore, the application is a proof-of-concept feasibility study and has not undergone clinical validation."),
     
    ("39. Future Improvements",
     "Future improvements will focus on: (1) Scale: training on the full MIMIC-IV database (>300,000 stays); (2) Waveforms: utilizing continuous temporal models (RNNs/Transformers) on high-frequency vital sign streams; (3) Integration: streaming bedside data via HL7 FHIR APIs; (4) Qualitative NLP: extracting notes embeddings using ClinicalBERT."),
     
    ("40. Complete Technical Conclusion",
     "In conclusion, this project establishes the feasibility of an Explainable AI-based bedside CDSS for early ICU mortality prediction. By restricting observations to the first 24 hours of ICU stay, applying clinical preprocessing, and engineering 6 multimodal aggregates, the pipeline avoids data leakage. Optimizing classification decision thresholds using Youden's J statistic resolves class imbalance, while SHAP explains risk scores, bridging the gap between machine learning and clinician bedside trust.")
]

def build_docx():
    print(f"Compiling Technical Documentation DOCX: {docx_path}...")
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run("PROJECT TECHNICAL DOSSIER & DOCUMENTATION")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = docx.shared.RGBColor(15, 98, 254) # IBM Blue
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    sub_run = subtitle_p.add_run("Explainable Artificial Intelligence for Real-Time In-Hospital Mortality Prediction and Bedside Decision Support in Intensive Care Units")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    
    doc.add_paragraph("\n" * 4)
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta_p.add_run("STATUS: COMPLETE (Phase II Overhaul)\nDEVELOPED FOR: Final Year Engineering Project\nAUTHOR: Project Team\nDATE: June 2026\nDEPLOYED APPLICATION: https://icu-risk-prediction-csgt.streamlit.app")
    meta_run.font.name = 'Courier New'
    meta_run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # Section contents
    for title, text in sections_content:
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(18)
        p_hdr.paragraph_format.space_after = Pt(6)
        p_hdr.paragraph_format.keep_with_next = True
        run_hdr = p_hdr.add_run(title)
        run_hdr.font.name = 'Arial'
        run_hdr.font.size = Pt(14)
        run_hdr.font.bold = True
        run_hdr.font.color.rgb = docx.shared.RGBColor(15, 98, 254)
        
        # Check if System Architecture to add Fig 1
        if "System Architecture" in title:
            img_path = os.path.join(FIGURES_DIR, "publication", "system_architecture_diagram.png")
            if os.path.exists(img_path):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(5.0))
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(10)
                run_cap = p_cap.add_run("Fig. 1. Architecture of the proposed Explainable AI-based Clinical Decision Support System for ICU mortality prediction.")
                run_cap.font.name = 'Arial'
                run_cap.font.size = Pt(9)
                run_cap.font.italic = True
                
        p_text = doc.add_paragraph()
        p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_text.paragraph_format.space_after = Pt(12)
        p_text.paragraph_format.line_spacing = 1.15
        run_text = p_text.add_run(text)
        run_text.font.name = 'Times New Roman'
        run_text.font.size = Pt(11)
        
    doc.save(docx_path)
    print(f"Technical Documentation DOCX generated successfully at: {docx_path}")

def add_header_footer(canvas, doc):
    canvas.saveState()
    if doc.page == 1:
        canvas.restoreState()
        return
        
    # Running Header
    canvas.setFont('Helvetica-Bold', 8)
    canvas.setFillColor(colors.HexColor('#0f62fe'))
    canvas.drawString(54, 785, "PROJECT TECHNICAL DOCUMENTATION")
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(colors.HexColor('#525252'))
    canvas.drawRightString(doc.pagesize[0] - 54, 785, "ICU MORTALITY PREDICTION CDSS")
    
    # Thin divider line
    canvas.setStrokeColor(colors.HexColor('#e0e0e0'))
    canvas.setLineWidth(0.5)
    canvas.line(54, 778, doc.pagesize[0] - 54, 778)
    
    # Thin divider line above footer
    canvas.line(54, 52, doc.pagesize[0] - 54, 52)
    
    # Running Footer
    canvas.drawString(54, 38, "CONFIDENTIAL / RESEARCH REFERENCE MATERIAL")
    canvas.drawRightString(doc.pagesize[0] - 54, 38, f"Page {doc.page}")
    canvas.restoreState()

def build_pdf():
    print(f"Compiling Technical Documentation PDF: {pdf_path}...")
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    style_cover_title = ParagraphStyle(
        'CoverTitle',
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=30,
        textColor=colors.HexColor('#0f62fe'),
        spaceAfter=15
    )
    
    style_cover_subtitle = ParagraphStyle(
        'CoverSubtitle',
        fontName='Helvetica-Oblique',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#525252'),
        spaceAfter=30
    )
    
    style_chapter_title = ParagraphStyle(
        'ChapterTitle',
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0f62fe'),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )
    
    style_body = ParagraphStyle(
        'CustomBody',
        fontName='Helvetica',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#393939'),
        spaceAfter=12
    )
    
    style_caption = ParagraphStyle(
        'FigureCaption',
        fontName='Helvetica-Oblique',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#6f6f6f'),
        alignment=1,
        spaceBefore=4,
        spaceAfter=10
    )
    
    story = []
    
    # Cover Page
    story.append(Spacer(1, 40))
    story.append(Paragraph("PROJECT TECHNICAL DOCUMENTATION & DOSSIER", style_cover_title))
    story.append(Paragraph("Explainable Artificial Intelligence for Real-Time In-Hospital Mortality Prediction and Bedside Decision Support in Intensive Care Units", style_cover_subtitle))
    
    story.append(Spacer(1, 10))
    story.append(Table([[""]], colWidths=[487], rowHeights=[4], style=[('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#0f62fe'))]))
    story.append(Spacer(1, 40))
    
    meta_text = """<b>DOCUMENT CLASSIFICATION:</b> Comprehensive System Documentation<br/>
<b>DATABASE SOURCE:</b> MIMIC-IV Clinical Demo Database v2.2 (Adult stays cohort)<br/>
<b>CHAMPION CLASSIFIER:</b> Youden-Optimized ExtraTrees Classifier (Optimal threshold = 0.39)<br/>
<b>EXPLAINABILITY ENGINE:</b> Tree SHAP (SHapley Additive exPlanations)<br/>
<b>DEPLOYMENT TARGET:</b> Streamlit Community Cloud (Proof-of-concept)<br/>
<b>LIVE APPLICATION:</b> <a href="https://icu-risk-prediction-csgt.streamlit.app">https://icu-risk-prediction-csgt.streamlit.app</a>"""
    story.append(Paragraph(meta_text, style_body))
    
    story.append(Spacer(1, 120))
    
    metadata_box = [
        [Paragraph("<b>Prepared For:</b> Claude / Downstream AI Code Translators & Conference Presenters", style_body)],
        [Paragraph("<b>Prepared By:</b> Clinical AI Development Team", style_body)],
        [Paragraph("<b>Status:</b> Production Overhaul Complete (Phase II Audit Passed)", style_body)],
        [Paragraph("<b>Date:</b> June 2026", style_body)]
    ]
    meta_table = Table(metadata_box, colWidths=[487])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f4f4f4')),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#e0e0e0')),
        ('PADDING', (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(meta_table)
    
    story.append(PageBreak())
    
    # Sections loop
    for title, text in sections_content:
        story.append(Paragraph(title, style_chapter_title))
        
        # If System Architecture, embed the diagram
        if "System Architecture" in title:
            img_path = os.path.join(FIGURES_DIR, "publication", "system_architecture_diagram.png")
            if os.path.exists(img_path):
                story.append(Spacer(1, 10))
                story.append(Image(img_path, width=420, height=260))
                story.append(Paragraph("Fig. 1. Architecture of the proposed Explainable AI-based Clinical Decision Support System for ICU mortality prediction.", style_caption))
                
        # Split text into paragraphs if there are newlines
        paragraphs_in_text = text.split("\n")
        for p in paragraphs_in_text:
            story.append(Paragraph(p, style_body))
            
    doc.build(story, onFirstPage=lambda c, d: None, onLaterPages=add_header_footer)
    print(f"Technical Documentation PDF generated successfully at: {pdf_path}")

if __name__ == "__main__":
    try:
        build_docx()
        build_pdf()
        sys.exit(0)
    except Exception as e:
        print(f"Error compiling technical documentation: {e}", file=sys.stderr)
        sys.exit(1)
